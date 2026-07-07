from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.25

# 标题样式
for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.name = '微软雅黑'
    h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

# ── 封面 / 标题 ──
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.space_after = Pt(4)
run = title.add_run('Unity 3D 生存射击游戏')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.space_after = Pt(2)
run = sub.add_run('实习面试 · 项目经历 + 面试问答准备')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.space_after = Pt(18)
run = meta.add_run('开发周期：2026.06 – 2026.07  |  独立开发  |  Unity 2022 + C#')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ═══════════════════════════════════════
# PART 1 — 项目经历
# ═══════════════════════════════════════
doc.add_heading('第一部分：项目经历', level=1)

# --- 1.1 项目概述 ---
doc.add_heading('1.1  项目概述', level=2)
doc.add_paragraph(
    '基于 Unity 2022 Built-in Render Pipeline 开发的一款第三人称生存射击游戏。'
    '玩家通过鼠标瞄准、键盘移动，在封闭场景中使用枪械射击持续生成的敌人，'
    '击杀敌人获得分数，玩家阵亡后可重开游戏。项目覆盖角色控制、射击系统、'
    '敌人 AI、血量管理、伤害反馈、分数系统、摄像机跟随等完整玩法闭环。'
)

# --- 1.2 技术栈 ---
doc.add_heading('1.2  技术栈', level=2)
table = doc.add_table(rows=8, cols=2, style='Light Grid Accent 1')
data = [
    ('引擎', 'Unity 2022（Built-in Render Pipeline）'),
    ('语言', 'C#'),
    ('物理系统', 'Rigidbody、CapsuleCollider、SphereCollider（Trigger）'),
    ('导航系统', 'NavMesh + NavMeshAgent 自动寻路'),
    ('渲染 & 特效', 'LineRenderer（弹道线）、ParticleSystem（枪口粒子 / 命中特效）、Light（枪口闪光）'),
    ('动画系统', 'Animator 状态机（Idle → Walk → Death）、Animation Event'),
    ('UI 系统', 'Unity UI（Text 血量/分数显示、Image 受伤红色闪烁）'),
    ('编辑器扩展', 'MenuItem 自定义工具、AssetDatabase 批量资源处理'),
]
for i, (k, v) in enumerate(data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v
    for cell in table.rows[i].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9.5)

doc.add_paragraph()

# --- 1.3 核心模块 ---
doc.add_heading('1.3  核心模块与实现细节', level=2)

modules = [
    ('模块一：玩家移动与朝向（MyPlayerMovement.cs）',
     [
         '采用 Update 读输入 + FixedUpdate 做物理移动的双帧分离架构，'
         '避免 Input 数据丢失和物理模拟抖动',
         '鼠标位置通过 Camera.main.ScreenPointToRay 发射射线检测 Floor 图层（LayerMask），'
         '获取世界坐标交点，计算 Quaternion.LookRotation 后用 Rigidbody.MoveRotation 实现朝向',
         '移动向量归一化后乘以 Time.fixedDeltaTime，保证不同帧率下移动速度一致',
     ]),
    ('模块二：枪械射击系统（MyPlayerShooting.cs）',
     [
         '枪口沿 transform.forward 方向发射 Physics.Raycast，通过 LayerMask 仅检测 Enemy 图层，'
         '避免误射地面或自身',
         '命中后通过 GetComponent<MyEnemyHealth>() 获取敌方血量组件，'
         '调用 TakeDamage(damage, hitPoint) 传递伤害值和命中世界坐标',
         'LineRenderer 绘制弹道轨迹：SetPosition(0, 枪口) → SetPosition(1, 命中点或最大射程)',
         '开火间隔用 timeBetween（0.15s）控制射速，枪口闪光和弹道线通过 efforts 系数'
         '提前熄灭（0.03s），模拟真实枪械瞬发闪光效果',
         '所有组件（GunAudio、GunLight、Gunline、GunParticle）在 Awake 中缓存引用，'
         '调用前均做 null 检查，防止组件缺失导致静默崩溃',
     ]),
    ('模块三：敌人 AI 系统（MyEnemyMovement.cs + MyEnemyAttack.cs）',
     [
         '寻路：NavMeshAgent.SetDestination 持续追踪玩家 Transform.position，'
         '在自身存活（health > 0）且玩家未死亡时执行，否则 nav.enabled = false',
         '攻击检测：CapsuleCollider（IsTrigger）检测敌人与玩家的距离，'
         'OnTriggerEnter 设 PlayerInRange = true，OnTriggerExit 设 false',
         'Update 中以 0.8s 间隔（time >= 0.8f）循环调用 Attack()，'
         '通过 GetComponent 获取玩家血量组件并调用 TakeDamage',
     ]),
    ('模块四：伤害与死亡系统（MyEnemyHealth.cs + MyPlayerHealth.cs）',
     [
         '敌人受伤：TakeDamage(damage, hitPoint) 接收伤害值和命中世界坐标，'
         '在命中点播放 ParticleSystem 粒子特效 + 受击音效，血量归零触发 Death()',
         '敌人死亡：禁用 NavMeshAgent 和 CapsuleCollider（防止继续寻路和被重复命中），'
         '播放 Death 动画，切换死亡音效。通过 Animation Event 回调 StartSinking() '
         '实现尸体下沉动画 + 2 秒后 Destroy(gameObject)',
         '敌人死亡时通过静态变量 MyPlayerScores.Scores += 10 累加分数',
         '玩家受伤：扣血 → 更新 UI Text → Image 红色闪烁（FlashColor + Lerp 渐变恢复）',
         '玩家死亡：isPlayerDead 标记 → 禁用移动/射击组件 → 播放 Death 动画 → '
         '通知敌人 Animator 触发 PlayerDeath → 播放死亡音效 → Rigidbody.isKinematic = true 冻结',
     ]),
    ('模块五：敌人管理器（MyEnemyManager.cs）',
     [
         '使用 InvokeRepeating("Spawn", firstTime, interval) 定时在指定生成点 '
         'Instantiate 敌人 Prefab，可在 Inspector 面板调节生成间隔和首次生成延迟',
     ]),
    ('模块六：摄像机跟随（CameraFllow.cs）',
     [
         'Start 时计算摄像机与玩家的初始偏移量 offset = cameraPos - playerPos',
         'FixedUpdate 中用 Vector3.Lerp 从当前位置向 targetPos 平滑插值，'
         'Smoothing 系数通过 public 变量在 Inspector 中调节',
     ]),
    ('模块七：编辑器工具（FixMaterialsEditor.cs）',
     [
         '使用 [MenuItem("Tools/...")] 注册自定义菜单项到 Unity Editor',
         '遍历 Assets/Materials/ 目录下所有 .mat 文件，检测 Shader 状态',
         '若 Shader 为 Hidden / Error / null，自动替换为 Standard Shader',
         '兼容 Built-in / URP / HDRP 三套渲染管线的 Shader 查找（Shader.Find 回退机制）',
         '使用 EditorUtility.DisplayDialog 弹窗反馈修复结果',
         '通过 AssetDatabase.SaveAssets + Refresh 确保修改持久化',
     ]),
]

for title_text, bullets in modules:
    doc.add_heading(title_text, level=3)
    for b in bullets:
        p = doc.add_paragraph(b, style='List Bullet')

# --- 1.4 问题排查 ---
doc.add_heading('1.4  典型问题排查与解决', level=2)

table2 = doc.add_table(rows=7, cols=4, style='Light Grid Accent 1')
headers = ['问题现象', '排查方法', '根因', '解决方案']
for i, h in enumerate(headers):
    table2.rows[0].cells[i].text = h
    for p in table2.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)

problems = [
    ('弹道指向世界原点 (0,0,0)',
     'Debug.Log 打印 LineRenderer 的 position 坐标，发现数值正确但渲染错误',
     'LineRenderer.useWorldSpace = False，世界坐标被当成局部偏移解释',
     '在 Shoot() 中显式设置 Gunline.useWorldSpace = true'),
    ('敌人命中后不掉血',
     '逐层加 Log：Shoot() → Raycast → GetComponent → TakeDamage 全链路排查',
     'ParticleSystem 组件缺失导致 Shoot() 中途崩溃，后续伤害代码不执行',
     '补全组件 + 所有组件调用添加 null 检查'),
    ('敌人模型显示为纯白色，无纹理',
     '对比 FBX 的 .meta 文件 externalObjects 映射表与 Materials/*.mat 的 GUID',
     'FBX 材质重映射错乱，身体材质指向了 EyesMaterial（无 MainTex 贴图）',
     'Inspector → Materials 标签页修正 Remapped Materials 下拉框'),
    ('玩家死亡音效重复播放',
     '排查 Death() 和 TakeDamage() 的完整调用链',
     'TakeDamage 中 PlayerAudioSource.Play() 与 Death() 中 Play() 重复触发',
     '区分受伤音效和死亡音效，死亡时先赋值 clip 再 Play'),
    ('敌人尸体不消失',
     '对比参考 EnemyHealth 发现 Death() 末尾调用了 StartSinking()',
     '自定义 MyEnemyHealth.Death() 遗漏调用 StartSinking()',
     '在 Death() 末尾补充调用，通过 Animation Event 触发'),
    ('部分材质 Shader 显示为 Hidden',
     '检查材质 Inspector 面板，发现 Shader 引用丢失',
     '项目使用 Built-in RP 但部分材质引用了不存在的 Shader',
     '编写 Editor 工具批量修复所有材质的 Shader 引用'),
]

for i, (a, b, c, d) in enumerate(problems):
    row = table2.rows[i + 1]
    row.cells[0].text = a
    row.cells[1].text = b
    row.cells[2].text = c
    row.cells[3].text = d
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

doc.add_paragraph()

# --- 1.5 项目收获 ---
doc.add_heading('1.5  项目收获', level=2)

gains = [
    '组件化设计思想：伤害系统通过 GetComponent<T>() 实现松耦合——射击脚本不依赖具体敌人类型，任何挂载了 MyEnemyHealth 的对象均可被命中扣血',
    'Unity 帧循环理解：掌握了 Update（读输入/动画）与 FixedUpdate（物理操作）的调用时序差异及最佳实践',
    '射线检测完整链路：ScreenPointToRay → LayerMask → Raycast → RaycastHit，构建了 Unity 3D 交互的核心技术认知',
    'Unity 序列化机制：通过 .meta 文件的 GUID 引用链和 FBX 的 externalObjects 映射，深入理解了资源引用与重映射原理',
    '调试方法论：复杂交互链路（射击→命中→扣血→死亡）排查时，逐层加 Debug.Log，不跳过任何环节',
    '编辑器扩展能力：能够编写 Editor 脚本自动化处理批量资源问题，提升开发效率',
    '代码健壮性意识：养成对 GetComponent、FindGameObjectWithTag 等 API 返回值做 null 检查的习惯，防止静默崩溃',
]
for g in gains:
    doc.add_paragraph(g, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════
# PART 2 — 面试问答
# ═══════════════════════════════════════
doc.add_heading('第二部分：面试官可能问的问题', level=1)

doc.add_paragraph(
    '以下问题基于本项目的技术栈和实现细节整理，按考察维度分类。'
    '每个问题后面附了参考答案要点，面试时用自己的话组织即可。'
)

# --- 2.1 C# 与 Unity 基础 ---
doc.add_heading('2.1  C# 与 Unity 基础', level=2)

qa_basic = [
    ('Q1：Update 和 FixedUpdate 的区别？为什么移动要放 FixedUpdate？',
     [
         'Update 每帧调用一次，调用频率不固定（取决于帧率），适合读取输入、更新动画参数',
         'FixedUpdate 按固定时间步长调用（默认 0.02s = 50次/秒），与物理引擎同步',
         'Rigidbody 的 MovePosition / MoveRotation 属于物理操作，放在 FixedUpdate 中能保证物理模拟稳定，避免因帧率波动导致移动忽快忽慢',
         '我的做法：Update 中通过 Input.GetAxisRaw 读取输入并存储到成员变量，FixedUpdate 中消费这些变量做物理移动',
     ]),
    ('Q2：GetComponent 的性能开销大吗？你是怎么用的？',
     [
         'GetComponent 有一定性能开销，尤其是在 Update 中频繁调用',
         '我的做法：在 Awake() 中把所有需要的组件引用缓存到成员变量（如 GunAudio、Gunline、GunParticle），后续直接使用缓存引用',
         '射击时用 GetComponent<MyEnemyHealth>() 获取命中对象的血量组件，这是射线命中时的一次性操作，开销可接受',
         '所有 GetComponent 返回值都做了 null 检查，防止组件缺失导致 NullReferenceException',
     ]),
    ('Q3：LayerMask 是什么？你在项目中怎么用的？',
     [
         'LayerMask 是 Unity 的层级过滤机制，用于选择性检测碰撞/射线',
         '我的射击系统用 LayerMask.GetMask("Enemy") 让射线只检测 Enemy 图层，避免打到地面、墙壁或玩家自身',
         '玩家朝向系统用 LayerMask.GetMask("Floor") 让射线只检测地面，确保鼠标指向的是有效的地面位置',
         '好处：减少不必要的碰撞检测，提升性能，同时业务逻辑更清晰',
     ]),
    ('Q4：Animator 状态机你是怎么配置的？Animation Event 做了什么？',
     [
         'Animator 配置了 Idle → Walk → Death 的状态转换，通过 SetBool("IsWalking") 和 SetTrigger("Death") 控制',
         'Animation Event 是在动画片段的时间轴上挂载回调函数，当动画播放到指定帧时自动调用',
         '我在敌人 Death 动画的末尾帧挂载了 StartSinking() 方法，实现尸体下沉动画 + 延迟 Destroy，避免在动画中间突然删除对象',
     ]),
    ('Q5：public 和 [SerializeField] private 有什么区别？你用哪种？',
     [
         'public 变量在 Inspector 中可见且可被其他脚本直接访问，破坏了封装性',
         '[SerializeField] private 变量仅在 Inspector 中可见，外部脚本无法访问，封装性更好',
         '我在项目中主要用了 public（如 damage、Speed），因为当时为了方便调试；但如果重构，会用 [SerializeField] private 配合 public 属性',
     ]),
]

for q, answers in qa_basic:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    run.font.size = Pt(10.5)
    for a in answers:
        doc.add_paragraph(a, style='List Bullet')

# --- 2.2 物理与碰撞 ---
doc.add_heading('2.2  物理与碰撞检测', level=2)

qa_physics = [
    ('Q6：Raycast 的原理是什么？你的射击系统用了几种射线？',
     [
         'Raycast 从起点沿方向发射一条不可见的射线，检测途中碰到的 Collider，返回是否命中及命中信息（RaycastHit）',
         'Physics.Raycast 的参数包括：起点、方向、最大距离、LayerMask，还可以用 out 参数获取详细命中信息',
         '项目中用了两类射线：',
         '  ① 射击射线：Physics.Raycast(shootRay, out ShootHit, 100, ShootMask)，检测敌人',
         '  ② 鼠标射线：Physics.Raycast(cameraRay, out floorHit, 100, floorLayer)，检测地面用于玩家朝向',
         '命中后从 ShootHit 可获取：命中点坐标（point）、被命中的 Collider、碰撞法线等',
     ]),
    ('Q7：Trigger 和普通 Collider 的区别？你项目中怎么用的？',
     [
         '普通 Collider：产生物理碰撞响应（弹开、阻挡），触发 OnCollisionEnter/Exit/Stay',
         'Trigger（IsTrigger = true）：不产生物理碰撞，仅触发 OnTriggerEnter/Exit/Stay，物体可以穿过',
         '我的敌人攻击系统：CapsuleCollider 设为 Trigger，用于检测玩家是否进入攻击范围',
         '选 Trigger 的原因：敌人需要能穿过玩家继续移动（NavMeshAgent 寻路），不能因为碰撞把玩家推开',
         '玩家的 CapsuleCollider 不是 Trigger，用于正常的物理碰撞和受击检测',
     ]),
    ('Q8：Rigidbody 的 isKinematic 你什么时候设为 true？',
     [
         'isKinematic = true 时，Rigidbody 不受物理引擎的力和碰撞影响，但可以通过代码控制位置',
         '玩家死亡时设为 true：防止尸体被后续物理碰撞推动，保持在死亡位置',
         '敌人死亡时设为 true：配合禁用 NavMeshAgent，让尸体停在原地等待下沉动画',
     ]),
]

for q, answers in qa_physics:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    run.font.size = Pt(10.5)
    for a in answers:
        doc.add_paragraph(a, style='List Bullet')

# --- 2.3 架构与设计 ---
doc.add_heading('2.3  架构与设计模式', level=2)

qa_arch = [
    ('Q9：你的伤害系统是怎么设计的？为什么用 GetComponent 而不是直接引用？',
     [
         '伤害系统采用组件化设计：Shooting 脚本通过射线检测命中对象，用 GetComponent<MyEnemyHealth>() 获取血量组件，然后调用 TakeDamage()',
         '这是一种松耦合设计：射击系统不依赖具体的敌人类型，只要对象挂载了 MyEnemyHealth 组件就能被伤害',
         '优势：方便扩展——如果以后加 Boss 敌人、可破坏物体、队友误伤等，只需给它们挂载对应的血量组件即可，射击代码不需要修改',
         'TakeDamage 方法的参数设计也考虑了扩展性：传递伤害值 + 命中世界坐标，血量组件自己决定如何响应',
     ]),
    ('Q10：如果让你重构这个项目，你会怎么改进？',
     [
         '① 引入事件系统（UnityEvent 或 C# event）：用 OnEnemyDeath、OnPlayerDeath 等事件替代直接调用，进一步解耦各模块',
         '② 用 ScriptableObject 做数据配置：将敌人血量、伤害值、移动速度等配置抽离到 ScriptableObject 中，策划可以直接在 Inspector 中调数值而不用改代码',
         '③ 对象池（ObjectPool）：敌人频繁生成销毁，用对象池替代 Instantiate/Destroy 可减少 GC 和性能开销',
         '④ 状态管理：引入有限状态机（FSM）管理玩家和敌人的状态切换（Idle、Walk、Attack、Death），替代 bool 变量堆叠',
         '⑤ 输入系统升级：用 Unity New Input System 替代旧版 Input Manager，支持多设备和按键映射',
         '⑥ 用 [SerializeField] private 替代 public 变量，提高封装性',
     ]),
    ('Q11：你提到了 Animation Event 触发 StartSinking，为什么不用代码直接调用？',
     [
         'Animation Event 保证 StartSinking() 在死亡动画播放到指定帧时才触发，不会在动画刚开始就下沉',
         '如果直接在 Death() 中调用，尸体可能在倒地动画播放过程中就开始下沉，视觉上不协调',
         '这是一种"动画驱动逻辑"的方式，让动画师（或你自己配置的动画）控制逻辑时机，减少代码中的时间耦合',
         'Destroy 的 2s 延迟则给尸体下沉动画留出播放时间',
     ]),
    ('Q12：InvokeRepeating 和协程（Coroutine）的区别？为什么用 InvokeRepeating？',
     [
         'InvokeRepeating：按固定时间间隔重复调用指定方法，简单直接',
         'Coroutine：更灵活，可以 yield return new WaitForSeconds 控制等待，可以在循环中根据条件动态调整间隔',
         '敌人管理器场景中不需要动态调整生成间隔，InvokeRepeating 足够简洁',
         '如果要实现难度递增（生成越来越快），协程会更合适：yield return new WaitForSeconds(interval) 然后在循环中递减 interval',
     ]),
]

for q, answers in qa_arch:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    run.font.size = Pt(10.5)
    for a in answers:
        doc.add_paragraph(a, style='List Bullet')

# --- 2.4 渲染与特效 ---
doc.add_heading('2.4  渲染与特效', level=2)

qa_render = [
    ('Q13：LineRenderer 的 useWorldSpace 是什么意思？你遇到了什么问题？',
     [
         'useWorldSpace = true：坐标按世界空间解释，(0,0,0) 就是世界原点',
         'useWorldSpace = false：坐标按局部空间解释（相对于挂载对象的 Transform），(0,0,0) 就是对象自身位置',
         '我遇到的问题是：Inspector 中 useWorldSpace 被意外设为 false，但代码传入的是世界坐标（枪口世界坐标和命中点世界坐标），导致弹道终点严重偏移',
         '这就是典型的"坐标系不一致"问题，解决方法是代码中显式设置 Gunline.useWorldSpace = true',
         '收获：不要依赖 Inspector 的默认值，关键属性在代码中显式设置',
     ]),
    ('Q14：ParticleSystem 你在哪些地方用了？',
     [
         '枪口粒子（GunParticle）：每次射击时 Play()，模拟枪口火焰',
         '敌人受击粒子（enemyParticle）：在命中位置（ShootHit.point）播放，模拟弹孔/血迹/火花',
         '都做了 null 检查，因为最初调试时发现 ParticleSystem 缺失会导致整个 Shoot() 崩溃',
     ]),
]

for q, answers in qa_render:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    run.font.size = Pt(10.5)
    for a in answers:
        doc.add_paragraph(a, style='List Bullet')

# --- 2.5 调试与问题排查 ---
doc.add_heading('2.5  调试与问题排查', level=2)

qa_debug = [
    ('Q15：说一个你印象最深的 Bug，你是怎么定位和解决的？',
     [
         '推荐回答"弹道指向世界原点"这个问题，逻辑链条最完整：',
         '① 现象：枪口射线始终指向 (0,0,0)，明显不正常',
         '② 初步排查：Debug.Log 打印射线的起点和方向，发现 transform.forward 方向正确',
         '③ 深入排查：怀疑 LineRenderer 渲染有问题，打印 SetPosition 传入的坐标值——数值正确，说明问题在渲染层',
         '④ 关键发现：打印 Gunline.useWorldSpace，发现值为 False',
         '⑤ 根因：世界坐标被 LineRenderer 当成局部坐标解释，(100, 0, 50) 表示从挂载对象偏移这么多，而不是世界空间中的那个点',
         '⑥ 修复：Gunline.useWorldSpace = true 一行代码解决',
         '⑦ 反思：这是典型的坐标系不一致问题——数据的坐标系和渲染器的坐标系不匹配',
     ]),
    ('Q16：你怎么排查"敌人不掉血"这种问题？',
     [
         '采用"逐层加 Log"的方法，把整个调用链拆成多个检查点：',
         '① Shoot() 是否被调用？→ Update 中加了时间判断，确认 Fire1 按下且计时器满足条件',
         '② Raycast 是否命中？→ 打印 ShootHit.collider.name，确认命中敌人',
         '③ GetComponent 是否成功？→ 打印 health != null，确认拿到了 MyEnemyHealth 组件',
         '④ TakeDamage 是否执行？→ 在 TakeDamage 第一行加 Log，发现不输出',
         '⑤ 进一步排查：发现 ParticleSystem 缺失，Shoot() 中 GunParticle.Play() 抛出异常导致方法中断，后面的伤害代码没执行',
         '⑥ 教训：代码中任何可能为 null 的组件引用都要检查，异常链路要完整保护',
     ]),
    ('Q17：FBX 材质显示纯白色，你是怎么排查的？',
     [
         '① 第一步：检查材质文件（.mat），确认 Shader 和贴图引用正常 → 材质文件没问题',
         '② 第二步：怀疑 FBX 导入时的材质重映射 → 打开 FBX 的 .meta 文件',
         '③ 第三步：对比 .meta 中 externalObjects 的 GUID 映射与 Materials 目录下 .mat 文件的 GUID',
         '④ 发现：HellephantMaterial 的 GUID 对应的是 20c2c86...，但 externalObjects 中身体材质位置引用的是 EyesMaterial 的 GUID',
         '⑤ 修复：在 Inspector → Materials 标签页修正 Remapped Materials 下拉框',
         '⑥ 收获：FBX 导入时 Unity 会自动创建材质映射，但这个映射可能出错，需要直接看 .meta 文件确认 GUID 引用链',
     ]),
]

for q, answers in qa_debug:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    run.font.size = Pt(10.5)
    for a in answers:
        doc.add_paragraph(a, style='List Bullet')

# --- 2.6 编辑器扩展 ---
doc.add_heading('2.6  编辑器扩展', level=2)

qa_editor = [
    ('Q18：你写的编辑器工具做了什么？怎么实现的？',
     [
         '写了一个 FixMaterialsEditor，通过菜单 Tools → Fix All Materials 一键修复所有材质的 Shader 问题',
         '实现步骤：',
         '  ① 用 Shader.Find 递归尝试 Standard → URP/Lit → HDRP/Lit，兼容不同渲染管线',
         '  ② 用 Directory.GetFiles 遍历 Materials 文件夹下所有 .mat 文件',
         '  ③ 通过 AssetDatabase.LoadAssetAtPath 加载每个材质，检查 shader.name 是否为 Hidden/Error/null',
         '  ④ 需要修复的材质：mat.shader = standardShader + EditorUtility.SetDirty',
         '  ⑤ 最后 AssetDatabase.SaveAssets + Refresh 确保持久化',
         '  ⑥ 用 EditorUtility.DisplayDialog 向用户弹窗反馈修复结果',
         '用的 API：MenuItem（注册菜单）、AssetDatabase（资源操作）、EditorUtility（弹窗/进度条）',
     ]),
]

for q, answers in qa_editor:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    run.font.size = Pt(10.5)
    for a in answers:
        doc.add_paragraph(a, style='List Bullet')

# --- 2.7 行为问题 ---
doc.add_heading('2.7  行为/综合问题', level=2)

qa_behavior = [
    ('Q19：这个项目你做的时候遇到的最大困难是什么？',
     [
         '最大困难是"弹道指向错误 + 敌人不掉血 + 材质不显示"三个问题叠加，导致整个战斗系统完全不工作',
         '困难在于：同时有多个问题，现象互相干扰，不知道先修哪个',
         '解决思路：从数据流源头开始排查——先确认射击是否触发 → 射线是否命中 → 命中信息是否正确 → 伤害是否传递 → 视觉是否正确',
         '每个环节用 Debug.Log 打桩，从上到下逐个排除，最终把三个问题全部定位',
         '最大的心得：面对多问题叠加时要冷静，逐个击破，不能同时改多个地方（可能引入新问题）',
     ]),
    ('Q20：这个项目你是怎么安排时间的？做了多久？',
     [
         '总共约 2-3 周，分为几个阶段：',
         '① 第 1 周：搭建基础框架——玩家移动、摄像机跟随、射击系统、敌人 AI 寻路',
         '② 第 2 周：完善战斗闭环——伤害系统、血量管理、分数系统、死亡逻辑',
         '③ 第 3 周：调试修复、视觉效果调优、写编辑器工具',
         '每个阶段先保证核心功能跑通，再逐步完善细节和边界情况处理',
     ]),
    ('Q21：这个项目你从中学到最多的是什么？',
     [
         '技术层面：对 Unity 的物理系统、射线检测、序列化机制有了深入理解',
         '方法论层面：学会了系统性地排查问题——Debug.Log 逐层打桩、从数据流源头追踪、不跳步',
         '工程习惯：养成了 null 检查、组件缓存、代码注释的习惯',
         '最大的改变：以前遇到 Bug 会乱试，现在会先分析再动手，效率提高了',
     ]),
]

for q, answers in qa_behavior:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    run.font.size = Pt(10.5)
    for a in answers:
        doc.add_paragraph(a, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════
# PART 3 — 面试自查清单
# ═══════════════════════════════════════
doc.add_heading('第三部分：面试前自查清单', level=1)

checklist = [
    '□ 能否用自己的话讲清楚项目整体架构？（30 秒版本 + 2 分钟版本各准备一份）',
    '□ 能否在白板上画出模块间的调用关系？（玩家 → 射击 → Raycast → 敌人血量 → 死亡 → 分数）',
    '□ 每个核心脚本的 Awake / Update / FixedUpdate 里做了什么，能不能不看代码说出来？',
    '□ LineRenderer.useWorldSpace = false 导致的问题能不能用一句话解释清楚？',
    '□ GetComponent 的性能开销和缓存策略能说出来吗？',
    '□ NavMeshAgent 的原理（烘焙 → 寻路 → SetDestination）能简单解释吗？',
    '□ Trigger vs Collider 的区别和使用场景能举出本项目中的例子吗？',
    '□ 如果要加一个"Boss 敌人"，需要改哪些代码？能不能不改射击代码就支持？',
    '□ Unity 的 .meta 文件和 GUID 是干什么的？FBX 的 externalObjects 是什么？',
    '□ 编辑器扩展的 MenuItem 和 AssetDatabase 的用法能简单描述吗？',
]

for item in checklist:
    doc.add_paragraph(item, style='List Bullet')

# ── 保存 ──
desktop = os.path.join(os.path.expanduser('~'), 'Desktop')
output_path = os.path.join(desktop, 'ShooterGame_面试准备.docx')
doc.save(output_path)
print(f'文档已保存至: {output_path}')
